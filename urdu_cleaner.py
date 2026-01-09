"""
URDU-ONLY OCR CLEANER & DOCX REBUILDER
------------------------------------
• Paragraph-level Urdu processing
• Nastaliq-safe layout
• Header detection & removal
• Urdu spell highlighting
• OCR normalization
• Soft-split paragraph merging
• NO English logic
"""

import re
import csv
from pathlib import Path
import docx
from collections import Counter
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from wordfreq import zipf_frequency
from docx.shared import Pt



# ============================================================
# CONSTANTS (HARD-LOCKED AS REQUESTED)
# ============================================================

URDU_FONT = "Jameel Noori Nastaleeq"
URDU_FONT_SIZE = Pt(14)
URDU_LANG = "ur-PK"
SPELL_THRESHOLD = 1.6   # aggressive (as requested)

URDU_CHAR_RE = re.compile(r'[\u0600-\u06FF\u0750-\u077F]')
URDU_WORD_RE = re.compile(r'[\u0600-\u06FF\u0750-\u077F]+')


# ============================================================
# BASIC UTILITIES
# ============================================================

def is_urdu_text(text: str) -> bool:
    return bool(URDU_CHAR_RE.search(text))

EN_WORD_RE = re.compile(r'^[\(\[\{"]*[A-Za-z]+[\)\]\}",\.\:;!?]*$')

def is_english_word(token: str) -> bool:
    return bool(EN_WORD_RE.match(token))

def normalize_urdu_text(text: str) -> str:
    """Unicode- and OCR-safe normalization for Urdu."""
    if not text:
        return ""

    # Normalize spaces
    text = text.replace("\u00A0", " ")
    text = re.sub(r"[ \t\f\v]+", " ", text)

    # Remove zero-width junk
    text = re.sub(r"[\u200B\u200C\u200D\uFEFF]", "", text)

    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Urdu punctuation normalization
    text = text.replace("٬", "،")
    text = text.replace("٫", "۔")

    return text.strip()


def is_valid_urdu_word(word: str) -> bool:
    """Frequency-based Urdu word validation."""
    word = re.sub(r"[\u200B\u200C\u200D\uFEFF]", "", word)
    return zipf_frequency(word, "ur") >= SPELL_THRESHOLD


# ============================================================
# HEADER DETECTION & REMOVAL (URDU-SAFE)
# ============================================================

AUDIT_HEADER = ["file", "para_index", "action", "pattern", "before_preview", "after_preview"]


def detect_repetitive_headers(paragraphs, min_repeats=3):
    normalized = []

    for p in paragraphs:
        line = p.strip()
        if not line:
            continue

        # Skip numeric-only or decorative lines
        if re.match(r'^[-–—]*\s*\(?\d{1,3}\)?\s*[-–—]*$', line):
            continue

        # Require some Urdu content
        if not is_urdu_text(line):
            continue

        normalized.append(line)

    counts = Counter(normalized)

    return {
        line for line, freq in counts.items()
        if freq >= min_repeats
    }


def remove_known_headers(paragraphs, detected_headers, audit_writer=None, filename=None):
    filtered = []
    seen = set()

    for i, para in enumerate(paragraphs):
        stripped = para.strip()

        if stripped in detected_headers:
            if stripped not in seen:
                seen.add(stripped)
                filtered.append(para)
            else:
                if audit_writer:
                    audit_writer.writerow([
                        filename or "",
                        i,
                        "remove_header",
                        "repetitive_header",
                        stripped[:80],
                        ""
                    ])
            continue

        filtered.append(para)

    return filtered

def clean_urdu_numeric_artifacts(text: str) -> str:
    """
    Urdu-safe numeric & artifact cleaner.
    Designed specifically for religious / academic Urdu OCR text.
    """

    if not text:
        return text

    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Split into paragraphs
    paras = text.split('\n\n')
    cleaned = []

    for para in paras:
        p = para.strip()

        if not p:
            continue

        # 1. Remove standalone numeric lines (Urdu or Latin digits)
        if re.fullmatch(r'[۰-۹0-9]{1,3}', p):
            continue

        # 2. Remove decorative page numbers: - 13 -, (14), — 15 —
        if re.fullmatch(r'[-–—]?\s*\(?[۰-۹0-9]{1,3}\)?\s*[-–—]?', p):
            continue

        # 3. Remove pure numeric footnotes like [12] or (12)
        if re.fullmatch(r'[\[\(]\s*[۰-۹0-9]{1,3}\s*[\]\)]', p):
            continue

        cleaned.append(p)

    text = '\n\n'.join(cleaned)

    # 5. Collapse excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()

# ============================================================
# URDU SPELL HIGHLIGHTING (SAFE)
# ============================================================

def highlight_urdu_misspellings(paragraph):
    """
    Highlight misspelled Urdu words.

    """
    text = paragraph.text
    if not text.strip():
        return

    misspelled = {
        w for w in URDU_WORD_RE.findall(text)
        if len(w) >= 3 and not is_valid_urdu_word(w)
    }

    if not misspelled:
        return

    # Clear paragraph text only (no formatting logic)
    paragraph._p.clear_content()

    parts = re.split(r'([\u0600-\u06FF\u0750-\u077F]+)', text)

    for part in parts:
        run = paragraph.add_run(part)

        # ONLY highlight — nothing else
        if part in misspelled:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def format_english_run(run):
    """
    Apply English (LTR) formatting to a RUN only.
    """

    rPr = run._element.get_or_add_rPr()

    # Explicit LTR (important in RTL documents)
    ltr = OxmlElement("w:ltr")
    rPr.append(ltr)

    # Language
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    rPr.append(lang)

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:fareast"), "Times New Roman")
    rPr.append(rFonts)

    # Style
    run.font.size = Pt(12)
    run.bold = False

# ============================================================
# PARAGRAPH FORMATTING (FINAL PASS)
# ============================================================

def set_run_urdu_properties(run, urdu_font_name = URDU_FONT, lang_code="ur-PK"):
    # 1) set font metadata (ASCII/complex script)
    try:
        run.font.name = urdu_font_name
        r = run._element
        # set complex script font (for Urdu)
        r.rPr.rFonts.set(qn('w:cs'), urdu_font_name)
        # also set ascii/hAnsi to the same font to help some viewers
        try:
            r.rPr.rFonts.set(qn('w:ascii'), urdu_font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), urdu_font_name)
        except Exception:
            pass
    except Exception:
        pass

    # 2) language tag
    try:
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        # remove existing w:lang if present
        for child in list(rPr):
            if child.tag == qn('w:lang'):
                rPr.remove(child)
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), lang_code)
        rPr.append(lang)
    except Exception:
        pass

    # 3) mark run as RTL at run-level (helps Word render & proofing)
    try:
        rtl_run = OxmlElement('w:rtl')
        rtl_run.set(qn('w:val'), "1")
        # remove existing run-level rtl if present
        for child in list(rPr):
            if child.tag == qn('w:rtl'):
                rPr.remove(child)
        rPr.append(rtl_run)
    except Exception:
        pass

def set_paragraph_rtl(paragraph):
    # 1) set paragraph alignment to justify (visual)
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        pass

    # 2) set paragraph-level RTL / bidi flags in XML
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    # remove existing flags if present, then add both to be safe
    try:
        # Remove any existing w:rtl or w:bidi children
        for child in list(pPr):
            if child.tag in (qn('w:rtl'), qn('w:bidi')):
                pPr.remove(child)
        # Add w:rtl w:val="1"
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), "1")
        pPr.append(rtl)
        # Also add w:bidi w:val="1" (some Word versions check this)
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), "1")
        pPr.append(bidi)
    except Exception:
        pass

def apply_urdu_font_size(run, size_pt):
    """
    Apply font size for complex scripts (Urdu).
    DOES NOT touch font name or other properties.
    """
    rPr = run._element.get_or_add_rPr()

    size_val = str(int(size_pt.pt * 2))  # Word uses half-points

    # Remove existing size nodes to avoid conflicts
    for child in list(rPr):
        if child.tag in (qn('w:sz'), qn('w:szCs')):
            rPr.remove(child)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), size_val)

    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), size_val)

    rPr.append(sz)
    rPr.append(szCs)



# ============================================================
# FINAL DOCX PROCESSOR (URDU-ONLY)
# ============================================================

def process_docx_file(input_path, output_docx, audit_csv_path=None, audit=False):
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    doc = docx.Document(str(input_path))
    raw_paras = [
        "".join(r.text for r in p.runs).strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]

    if audit:
        audit_file = open(audit_csv_path, "w", newline="", encoding="utf-8")
        audit_writer = csv.writer(audit_file)
        audit_writer.writerow(AUDIT_HEADER)
    else:
        audit_writer = None
        audit_file = None

    # Normalize OCR text (Urdu only)
    cleaned = [
        clean_urdu_numeric_artifacts(
            normalize_urdu_text(p)
        )
        for p in raw_paras
    ]

    # Header detection & removal
    headers = detect_repetitive_headers(cleaned)
    print(f"Detected {len(headers)} repetitive headers to remove.")
    cleaned = remove_known_headers(
        cleaned,
        headers,
        audit_writer=audit_writer,
        filename=input_path.name
    )

    # Build new document
    new = docx.Document()

    for para_text in cleaned:
        p = new.add_paragraph()
        p.add_run(para_text)

        # Highlight misspellings FIRST
        highlight_urdu_misspellings(p)

        # Final paragraph formatting
        # Final paragraph formatting (Urdu default)
        for run in p.runs:
            text = run.text.strip()
        
            if is_english_word(text):
                format_english_run(run)
            else:
                set_run_urdu_properties(run)
                apply_urdu_font_size(run, URDU_FONT_SIZE)

        # Paragraph direction (Urdu document default)
        set_paragraph_rtl(p)

    # Margins
    for sec in new.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    new.save(output_docx)

    if audit_file:
        audit_file.close()

    print("✅ Urdu document cleaned:", output_docx)


# ============================================================
# FOLDER PROCESSOR
# ============================================================

def process_folder(root_dir, output_dir, overwrite=False):
    root = Path(root_dir)
    output_root = Path(output_dir)

    if not root.exists():
        raise FileNotFoundError(root)

    for docx_file in root.rglob("*.docx"):
        if docx_file.name.startswith("~$"):
            continue
        if docx_file.stem.endswith(".cleaned"):
            continue

        rel_path = docx_file.relative_to(root)
        out_file = output_root / rel_path
        out_file = out_file.with_name(f"{out_file.stem}.cleaned.docx")

        out_file.parent.mkdir(parents=True, exist_ok=True)

        print(f"📂 Processing: {rel_path}")

        try:
            process_docx_file(
                input_path=str(docx_file),
                output_docx=str(out_file)
            )
        except Exception as e:
            print(f"❌ Failed: {docx_file} → {e}")

