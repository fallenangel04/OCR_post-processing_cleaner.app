"""
Robust OCR-cleaner 
- Conservative heuristics + audit CSV for every automatic correction.
- Jupyter-friendly: call the functions directly.
"""

import re
import csv
from pathlib import Path
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from wordfreq import zipf_frequency
class SafeDict:
    def __init__(self, threshold=2.5, custom_words=None):
        self.threshold = threshold
        self.custom_words = set(custom_words or [])

    def check(self, word):
        if not word:
            return False

        w = word.lower().strip(".,;:()[]{}")

        if w in self.custom_words:
            return True

        return zipf_frequency(w, "en") >= self.threshold

    def suggest(self, word):
        # Optional stub to avoid crashes
        return []


AUDIT_HEADER = ["file","para_index","action","pattern","before_preview","after_preview"]

# Utilities
EN_DICT = SafeDict(
    threshold=2.5,
    custom_words={
        "econometrics",
        "macroeconomics",
        "keynesian",
        "neoclassical",
    }
)

EN_WORD_RE = re.compile(r"^[A-Za-z][A-Za-z']+$")

def is_word(w, threshold=2.5):
    return zipf_frequency(w.lower(), "en") >= threshold

def is_valid_arabic(word):
    return zipf_frequency(word, "ar") >= 2.0

# def is_valid_urdu(word):
#     return zipf_frequency(word, "ur") >= 2.0

# URDU_WORD_RE = re.compile(r'^[\u0600-\u06FF\u0750-\u077F]+$')

# def is_urdu_word(token: str) -> bool:
#     return bool(URDU_WORD_RE.match(token))

AR_WORD_RE = re.compile(
    r'^[\u0600-\u06FF]+$'
)

URDU_ONLY_CHARS = set("ٹڈڑںھہے")

def is_arabic_word(token: str) -> bool:
    if not AR_WORD_RE.match(token):
        return False
    return not any(ch in URDU_ONLY_CHARS for ch in token)

def _safe_preview(s, n=140):
    return (s[:n] + "...") if len(s) > n else s
UNICODE_HYPHEN_PATTERN = re.compile(
    "[" +
    "\u2010"  # hyphen
    "\u2011"  # non-breaking hyphen
    "\u2012"  # figure dash
    "\u2013"  # en dash
    "\u2014"  # em dash
    "\u2015"  # horizontal bar
    "\u00AD"  # soft hyphen
    "]"
)

def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)

def highlight_misspellings_in_paragraph(paragraph):
    text = paragraph.text

    words = re.findall(r"\b[A-Za-z\u0600-\u06FF\u0750-\u077F']+\b", text)
    misspelled = set()

    for w in words:
        if len(w) < 4 or w.isupper():
            continue

        if EN_WORD_RE.match(w):
            if not is_word(w):
                misspelled.add(w)
        elif AR_WORD_RE.match(w):
            if not is_arabic_word(w):
                misspelled.add(w)

    if not misspelled:
        return

    clear_paragraph(paragraph)

    tokens = re.split(r"(\b[A-Za-z\u0600-\u06FF\u0750-\u077F']+\b)", text)

    for tok in tokens:
        run = paragraph.add_run(tok)
        if tok in misspelled:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def normalize_unicode_and_spaces(text: str) -> str:
    if text is None:
        return ""
    # Replace non-breaking spaces
    text = text.replace("\u00A0", " ").replace("\u2007", " ").replace("\u202F", " ")
    # Remove zero-width spaces
    text = re.sub(r"[\u200B\u200C\u200D\uFEFF]", "", text)
    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Preserve double newlines (paragraph separation)
    text = re.sub(r"[ \t\f\v]+", " ", text)  # collapse spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n",text)        # collapse 3+ newlines to 2
    return text.strip()

def normalize_em_dashes(text):
    """
    Normalize OCR dash artifacts into proper em dashes.
    Handles:
    - word - word
    - word--word
    - word-- word
    - word --word
    Avoids:
    - hyphenated words
    - numeric ranges
    - minus signs
    """

    
    # 1️⃣ Normalize double or more hyphens between letters → em dash
    text = re.sub(
        r'(?<=[A-Za-z])-{2,}(?=[A-Za-z])',
        '—',
        text
    )

    # 2️⃣ Normalize spaced dash between words → em dash
    text = re.sub(
        r'(?<=[A-Za-z])\s+[-–—]\s+(?=[A-Za-z])',
        ' — ',
        text
    )

    # 3️⃣ Normalize mixed spacing: word-- word / word --word
    text = re.sub(
        r'(?<=[A-Za-z])\s*-{2,}\s*(?=[A-Za-z])',
        '—',
        text
    )

    return text

# Common repetitive OCR headers (specific to this book or similar structures)
HEADER_PATTERNS = []

def detect_repetitive_headers(paragraphs, min_repeats=3):
    """
    Detect recurring short lines likely to be running headers or footers.
    Returns a set of repeated short lines.
    """
    normalized = []
    for p in paragraphs:
        line = p.strip()
        if not line:
            continue

        # Skip standalone page numbers or numbering artifacts
        if re.match(r"^[-–—]?\s*\(?\d{1,3}\)?\s*[-–—]?$", line):
            continue

        # Require at least 3 letters (skip roman numerals or very short tokens)
        if len(re.findall(r"[A-Za-z]", line)) <= 2:
            continue
        normalized.append(line)

    counts = Counter(normalized)

    repetitive = {
        line for line, freq in counts.items()
        if freq >= min_repeats and not line.islower() and not line[0].isdigit()
    }

    return repetitive

def remove_known_headers(
    paragraphs,
    detected_headers,
    header_patterns,
    audit_writer=None,
    filename=None
):
    filtered = []
    seen_headers = set()

    compiled_patterns = [re.compile(p) for p in header_patterns]

    for i, para in enumerate(paragraphs):
        stripped = para.strip()

        # 1️⃣ Regex-based header removal (static + dynamic)
        if any(pat.match(stripped) for pat in compiled_patterns):
            if stripped not in seen_headers:
                seen_headers.add(stripped)
                filtered.append(para)
            else:
                if audit_writer:
                    audit_writer.writerow([
                        filename or "",
                        i,
                        "remove_header",
                        "regex_header_duplicate",
                        _safe_preview(stripped),
                        ""
                    ])
            continue

        # 2️⃣ Exact detected-header removal (fallback)
        if stripped in detected_headers:
            if stripped not in seen_headers:
                seen_headers.add(stripped)
                filtered.append(para)
            else:
                if audit_writer:
                    audit_writer.writerow([
                        filename or "",
                        i,
                        "remove_header",
                        "exact_header_duplicate",
                        _safe_preview(stripped),
                        ""
                    ])
            continue

        filtered.append(para)

    return filtered



def fix_ocr_hyphenated_words(text):
    """
    Fix OCR-broken words:
    - Merge words broken with hyphen + space.
    - Remove hyphen if the merged word is a valid English word.
    - Keep hyphen if it forms a legitimate hyphenated word.
    """
    pattern = re.compile(r'([A-Za-z]+)-\s+([A-Za-z]+)')

    def merge_match(m):
        first, second = m.group(1), m.group(2)
        combined = first + second
        if is_word(combined.lower()):   # valid English word → remove hyphen
            return combined
        else:                           # not valid → keep hyphen
            return first + "-" + second

    for _ in range(3):  # repeat for multi-stage breaks
        text = pattern.sub(merge_match, text)
    return text



def enhanced_ocr_preclean(text):
    """Enhanced version of OCR cleaning specifically for academic texts like the Contemporary Economic Challenges"""
    if not text:
        return text

    #Normalize em_dash 
    text = normalize_em_dashes(text)
    # 1. Initial normalization
    text = normalize_unicode_and_spaces(text)

    # 2. Fix hyphenation and line breaks more aggressively
    text = fix_ocr_hyphenated_words(text)

    # 2. Normalize spaced hyphens
    text = re.sub(r'\s*-\s*', "-", text)
    
    # 3. Clean academic-specific markers
    text = re.sub(r'(?m)^\s*\[?[ivxIVX]+\]?\s*$','',text)  # Remove Roman numerals like [i], [ii]
    text = re.sub(r"(?m)^\*\s?.*\n?", "", text) # Remove any paragraph starting with an asterisk (*)
    text = re.sub(r'\s*\[\d{1,3}\]\s*', ' ', text)  # Remove [123]
    ##text = re.sub(r'\s?\(\s*[\d\s:;\-–—,]+\s*\)', '', text) # Remove (1, 2-3:4;5)

    # 4. Paragraph segmentation
    paras = text.split('\n\n')
    cleaned_paragraphs = []

    for i, para in enumerate(paras):
        text = para.strip()     

        # Skip empty lines
        if not text:
            continue

        # Skip empty or page-number-only lines
        if text.isdigit() or re.match(r'^[0-9]{1,3}$', text.strip()):
            continue

        # Skip small lines that are likely just page numbers like "- 3 -" or "(4)"
        if re.match(r'^[-–—]?\s*\(?\d{1,3}\)?\s*[-–—]?$', text.strip()):
            continue

        cleaned_paragraphs.append(text)


    # 5. Join cleaned paragraphs back
    text = '\n\n'.join(cleaned_paragraphs)

    # Protect time formats like 11.40 A.M or 9.15 P.M
    text = re.sub(r'\b(\d{1,2})\.(\d{2})\s*(A\.M|P\.M)\b',r'\1:\2 \3',text,flags=re.IGNORECASE)
    
    # 6. Remove numeric/footnote artifacts
    # Remove lines that START with digits followed by optional spaces and a letter,
    # then anything until end of line. Typically removes "12 Chapter …", "3A Stuff…"
    text = re.sub(r'^\d+\s*[A-Za-z].*?(?=\n|$)', '', text, flags=re.MULTILINE)
    # Remove 'Page 12' only if it's at the start of a line or alone
    text = re.sub(r'^(?:Page\s*\d+)\s*$', '', text, flags=re.MULTILINE)
    # Remove 1–2 digit numbers that appear DIRECTLY after a word starting with a letter:
    # Example: "word23" → "word"
    text = re.sub(r'(?<![A-Z])(?<=[a-z])\d+', '', text)
    # Remove lines starting with 1-2 digits followed by a word (no punctuation before)
    text = re.sub(r'^\s*\d{1,2}\s+[A-Za-z]+\b.*$', '', text, flags=re.MULTILINE)
    # Remove trailing blocks of lines that start with numbers — usually end-of-book index blocks.
    text = re.sub(r'(\n\s*\d{1,2}\s+.*)+\Z', '', text)
    # Remove (pg. 12) style footnotes
    text = re.sub(r'\(\s*pg\.?\s*\d+\s*\)', '', text, flags=re.IGNORECASE)
    # Replace dot followed by alphanumeric (".12", ".a", ".X12") with a plain dot.
    text = re.sub(r'(?<!\d)\.[0-9A-Za-z]+(?=\s|$)', '.', text)
    # Replace ".*" → "."
    text = re.sub(r'\.(\*)', '.', text)
    # Replace ". 12 Something…" → ". Something".
    text = re.sub(r'\.\s*\d+\s+(?=[A-Z])', '. ', text)
    # Remove numbers appearing right after an opening quotation mark (e.g., "23 → ")
    text = re.sub(r'"\s*(\d+)(?=\s|[A-Z])', '"', text)
    # Remove numbers immediately after a comma if followed by space and text or newline
    text = re.sub(r'(?<!\d),(\d{1,3})(?=\s+[A-Za-z\u0600-\u06FF])',',',text)
    # Remove numbers after a full stop followed by space if followed by capital letter or end of paragraph
    text = re.sub(r'\.\s+\d+(?=\s+[A-Z\u0600-\u06FF]|$)', '. ', text)
    # Remove digits that appear immediately after a colon or semicolon,
    # only if they are followed by whitespace, a paragraph break, or end of text
    text = re.sub(r'([:;])\d+(?=\s|$)', r'\1', text)


    # 7. Special character handling
    text = re.sub(r'["""]', '"', text)
    text = re.sub(r'[\u2018\u2019]', "'", text) # standardise quotation marks

    # 8. Final normalization
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'(?<=[.!?])\s{2,}(?=\S)', ' ', text)
    text = text.strip()

    return text


def merge_soft_split_paragraphs_text(text):
    """
    Merge soft-split paragraphs in a text string:
    - Current paragraph ends with lowercase letter and no punctuation
    - Next paragraph starts with lowercase letter
    """
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    merged_paras = []
    i = 0
    while i < len(paras):
        para = paras[i]

        # Merge soft split paragraphs
        while (
            i + 1 < len(paras) and
            para[-1].islower() and                  # current ends lowercase
            not para[-1] in ".!?" and               # current does not end with punctuation
            paras[i + 1].strip()[0].islower()       # next starts lowercase
        ):
            next_para = paras[i + 1].strip()
            para = para + " " + next_para
            i += 1  # skip merged paragraph

        merged_paras.append(re.sub(r"\s+", " ", para).strip())
        i += 1

    return "\n\n".join(merged_paras)

def merge_paragraphs_split_by_hyphen(text):
    """
    Merge paragraphs that are split by a line containing only a hyphen:
    - Keeps the hyphen in the merged sentence
    - Leaves other paragraphs untouched
    """
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    merged_paras = []
    i = 0

    while i < len(paras):
        para = paras[i]

        # Check if the next paragraph is a single hyphen
        while i + 1 < len(paras) and paras[i + 1].strip() == "-":
            next_para = paras[i + 2].strip() if i + 2 < len(paras) else ""
            # Merge current paragraph + hyphen + next paragraph
            para = f"{para} - {next_para}" if next_para else f"{para} -"
            i += 2  # skip the hyphen paragraph and the next merged paragraph

        merged_paras.append(re.sub(r"\s+", " ", para).strip())
        i += 1

    return "\n\n".join(merged_paras)

def merge_broken_words(paragraphs, audit_writer=None, filename=None):
    """
    Merge paragraphs split by broken hyphenated words,
    skipping numeric/page-number/junk paragraphs in between.
    """
    merged = []
    i = 0
    while i < len(paragraphs):
        cur = paragraphs[i].rstrip()
        # Check for hyphen at end of paragraph
        while True:
            m = re.search(r'([A-Za-z]+)-\s*$', cur)
            if not m:
                break

            # Look ahead for the next valid paragraph to merge
            j = i + 1
            next_para = ""
            while j < len(paragraphs):
                candidate = paragraphs[j].strip()
                # Skip junk paragraphs: page numbers, numeric, or single brackets
                if re.match(r'^(\[\d+\]|\d+|[-–—]+)$', candidate):
                    j += 1
                    continue
                next_para = candidate
                break

            if not next_para:
                break  # nothing to merge

            # Merge broken word + next paragraph
            full_word = m.group(1) + next_para.split()[0]
            rest = " ".join(next_para.split()[1:])
            cur = cur[:-(len(m.group(0)))] + full_word + " " + rest

            if audit_writer:
                audit_writer.writerow([
                    filename or "",
                    i,
                    "merge_broken_word_multi",
                    f"{m.group(1)}- + {next_para[:30]}...",
                    cur[:50]+"...",
                    cur[:50]+"..."
                ])

            # Move i forward to skip merged paragraphs
            i = j
        merged.append(cur)
        i += 1
    return merged

def format_arabic_run(run):
    """
    Apply Arabic formatting to a RUN only (Traditional Arabic font).
    """

    rPr = run._element.get_or_add_rPr()

    # RTL
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)

    # Language
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "ar-SA")
    rPr.append(lang)

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Traditional Arabic")
    rFonts.set(qn("w:hAnsi"), "Traditional Arabic")
    rFonts.set(qn("w:cs"), "Traditional Arabic")
    rFonts.set(qn("w:fareast"), "Traditional Arabic")
    rPr.append(rFonts)

    # Style

# def format_urdu_run(run):
#     """
#     Apply Urdu formatting to a RUN only.
#     """

#     rPr = run._element.get_or_add_rPr()

#     # RTL
#     rtl = OxmlElement("w:rtl")
#     rPr.append(rtl)

#     # Language
#     lang = OxmlElement("w:lang")
#     lang.set(qn("w:val"), "ur-PK")
#     rPr.append(lang)

#     # Font
#     rFonts = OxmlElement("w:rFonts")
#     rFonts.set(qn("w:ascii"), "Jameel Noori Nastaleeq")
#     rFonts.set(qn("w:hAnsi"), "Jameel Noori Nastaleeq")
#     rFonts.set(qn("w:cs"), "Jameel Noori Nastaleeq")
#     rFonts.set(qn("w:fareast"), "Jameel Noori Nastaleeq")
#     rPr.append(rFonts)

#     # Style
#     run.font.size = Pt(14)
#     run.bold = False

def ocr_preclean(text):
    """Apply OCR-specific cleaning before paragraph-level normalization."""
    text = enhanced_ocr_preclean(text)
    return text

def is_heading(text):
    return len(text.strip()) < 60

# -----------------------
# Full-document processing and I/O
# -----------------------
def process_docx_file(input_path, output_docx, audit_csv_path=None, audit=False):
    """
    Clean a single .docx file and write:
      - cleaned .docx (output_docx)
      - audit CSV listing automatic corrections (audit_csv_path)
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    # read docx paragraphs
    doc = docx.Document(str(input_path))
    # Merge runs inside each paragraph
    paras = ["".join(run.text for run in p.runs) for p in doc.paragraphs]

    # prepare audit CSV writer
    if audit:
        audit_rows_file = open(audit_csv_path, "w", newline="", encoding="utf-8")
        audit_writer = csv.writer(audit_rows_file)
        audit_writer.writerow(AUDIT_HEADER)
    else:
        class NullWriter:
            def writerow(self, *args, **kwargs):
                pass
        audit_writer = NullWriter()
        audit_rows_file = None
        

    # --- Detect repetitive headers automatically ---
    detected_headers = detect_repetitive_headers(paras, min_repeats=3)
    if detected_headers:
        print("\n📘 Detected possible repetitive headers:")
        for h in detected_headers:
            print("   •", h)
    # Clean each paragraph
    cleaned_paras = []
    for i, p in enumerate(paras):
        if not p or p.isspace():
            continue
        cleaned = ocr_preclean(p)
        if cleaned != p and audit:
            audit_writer.writerow([input_path.name, i, "ocr_preclean", "", _safe_preview(p), _safe_preview(cleaned)])
        cleaned_paras.append(cleaned)
        # Remove known repetitive headers (book title, Preface repeats, Appendix headers)
        # --- Dynamically remove detected repetitive headers ---
    auto_patterns = [re.escape(h) for h in detected_headers]
    
    # ALWAYS define local copy (no global mutation)
    local_header_patterns = HEADER_PATTERNS.copy()
    
    if auto_patterns:
        dynamic_header_regex = [fr'^\s*{p}\s*$' for p in auto_patterns]
        local_header_patterns.extend(dynamic_header_regex)
    
    cleaned_paras = remove_known_headers(
        cleaned_paras,
        detected_headers,
        header_patterns=local_header_patterns,
        audit_writer=audit_writer,
        filename=input_path.name
    )


    # Build new docx
    new = docx.Document()
    # --- Preserve original paragraph spacing
    cleaned_text = "\n\n".join(cleaned_paras)
    cleaned_text = merge_soft_split_paragraphs_text(cleaned_text)
    cleaned_text = merge_paragraphs_split_by_hyphen(cleaned_text)
    cleaned_paras = [p for p in cleaned_text.split("\n\n") if p.strip()]
    cleaned_paras = merge_broken_words(cleaned_paras, audit_writer=audit_writer, filename=input_path.name)

    for para_text in cleaned_paras:
        p = new.add_paragraph()

        # split into words + spaces (keeps punctuation intact)
        tokens = re.split(r"(\s+)", para_text)

        for token in tokens:
            if token.isspace():
                p.add_run(token)
                continue
        
            run = p.add_run(token)
        
            # if is_urdu_word(token):
            #     format_urdu_run(run)
        
            # if is_arabic_word(token):
            #     format_arabic_run(run)

        current_is_heading = is_heading(para_text)

        # Alignment logic
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Font logic
        if current_is_heading:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)

            
        # Update state
        highlight_misspellings_in_paragraph(p)
     
    for sec in new.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)
    new.save(output_docx)
    print("✅ Cleaned DOCX:", output_docx)

    if audit_rows_file:
        audit_rows_file.close()
        print("✅ Audit log:", audit_csv_path)

    return {
        "cleaned_docx": str(output_docx)
    }


# -----------------------
# Batch helper
# -----------------------
def process_folder(root_dir, output_dir, overwrite=False):
    root = Path(root_dir)
    output_root = Path(output_dir)

    if not root.exists():
        raise FileNotFoundError(root)

    for docx_file in root.rglob("*.docx"):
        if docx_file.name.startswith("~$"):
            continue
        if docx_file.stem.endswith(".cleaned"):
            continue

        rel_path = docx_file.relative_to(root)
        out_file = output_root / rel_path
        out_file = out_file.with_name(f"{out_file.stem}.cleaned.docx")

        out_file.parent.mkdir(parents=True, exist_ok=True)

        print(f"📂 Processing: {rel_path}")

        try:
            process_docx_file(
                input_path=str(docx_file),
                output_docx=str(out_file)
            )
        except Exception as e:
            print(f"❌ Failed: {docx_file} → {e}")

















